"""Document parsers + structural chunking (ADR-036, KB2a).

Pure, deterministic unit tests — no DB, no network, no model. Covers each format,
the named failure paths, and the bounded sentence-overlap chunker.
"""

from __future__ import annotations

import io

import pytest
from docx import Document
from pypdf import PdfWriter

from app.knowledge import chunk_document, estimate_tokens, parse_document
from app.knowledge.parsers import ParseError


def test_parse_txt() -> None:
    doc = parse_document(b"hello world\nsecond line", filename="a.txt")
    assert len(doc.sections) == 1
    assert doc.sections[0].heading_path is None
    assert doc.language == "en"


def test_parse_markdown_nested_headings() -> None:
    md = b"# Budget\n\n## 3.2 Approval\n\nSpend under 5w by dept lead.\n\n# Notes\n\nEnd."
    doc = parse_document(md, filename="b.md")
    paths = [s.heading_path for s in doc.sections]
    assert paths == ["Budget / 3.2 Approval", "Notes"]
    chunks = chunk_document(doc)
    assert [c.ordinal for c in chunks] == list(range(len(chunks)))
    assert chunks[0].heading_path == "Budget / 3.2 Approval"


def test_parse_docx_heading_and_cjk() -> None:
    d = Document()
    d.add_heading("季度预算说明", level=1)
    d.add_paragraph("单笔不超过5万由部门负责人审批，超过20万上报CFO。")
    buf = io.BytesIO()
    d.save(buf)
    doc = parse_document(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert doc.language == "zh"
    assert doc.sections[0].heading_path == "季度预算说明"
    assert "审批" in doc.sections[0].text


def test_parse_pdf_scanned_fails_named() -> None:
    w = PdfWriter()
    w.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    w.write(buf)
    with pytest.raises(ParseError) as ei:
        parse_document(buf.getvalue(), filename="scan.pdf")
    assert ei.value.code == "unsupported_scanned_pdf"


def test_unsupported_and_empty() -> None:
    with pytest.raises(ParseError) as e1:
        parse_document(b"\x00\x01\x02", filename="a.xyz")
    assert e1.value.code == "unsupported_type"
    with pytest.raises(ParseError) as e2:
        parse_document(b"   \n  ", filename="a.txt")
    assert e2.value.code == "empty_document"


def test_estimate_tokens_language_aware() -> None:
    # CJK ~ 1 token/char; latin ~ 1 token / 4 chars.
    assert estimate_tokens("预算审批") == 4
    assert estimate_tokens("abcdefgh") == 2


def test_chunker_splits_long_section_with_overlap() -> None:
    body = "".join(f"这是第{i}句测试内容。" for i in range(200))
    doc = parse_document(("# 长文\n\n" + body).encode("utf-8"), filename="c.md")
    chunks = chunk_document(doc, target_tokens=120, overlap_tokens=30)
    assert len(chunks) > 1
    assert [c.ordinal for c in chunks] == list(range(len(chunks)))
    # Every chunk keeps the section locator and stays near the target bound.
    assert all(c.heading_path == "长文" for c in chunks)
    assert all(c.token_estimate <= 120 + 40 for c in chunks)
    # Overlap: consecutive chunks share text (the carried tail).
    assert chunks[1].text[:6] in chunks[0].text
    # char_offset is non-decreasing within the single section.
    offsets = [c.char_offset for c in chunks]
    assert offsets == sorted(offsets)


def test_chunker_short_section_single_chunk() -> None:
    doc = parse_document(b"one short line", filename="s.txt")
    chunks = chunk_document(doc)
    assert len(chunks) == 1
    assert chunks[0].content_hash and len(chunks[0].content_hash) == 32
